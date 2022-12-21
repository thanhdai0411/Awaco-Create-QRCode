from docx import Document
from docx.shared import Inches
from docx.shared import Mm
import os
from time import sleep

import pyqrcode
import png
from pyqrcode import QRCode
from PIL import Image


document = Document()
p = document.add_paragraph()
r = p.add_run()

# set width page docx
section = document.sections[0]
# A3
section.page_height = Mm(65)  # 420  #280
section.page_width = Mm(280)  # 297  #65

section.left_margin = Mm(5)
section.right_margin = Mm(5)
section.top_margin = Mm(5)
section.bottom_margin = Mm(5)

section.header_distance = Mm(5)
section.footer_distance = Mm(5)

your_path = './QR'

qrData = []

print('Nhập số lượng QR code muốn in: ')
number_qr = input()

print('Nhập tiền tố bắt đầu: ')
prefix = input()

# print('Hello, ' + x)

for index in range(int(number_qr)):

    qrData.append(prefix)

for qr in qrData:
    url = pyqrcode.create(qr, error='H')
    # name = "cuon_tui_rac"
    print('Nhập tên QR Code: ')
    name = input()
    # https://app.awaco.vn/?gqr=id
    # 635a82a9144e8af5bd75eac7
    url.png('./QR/' + name + '.png', scale=6, quiet_zone=1)

    im = Image.open('./QR/' + name + '.png')
    im = im.convert("RGBA")
    im_size = im.resize((200, 200))

    a = Image.open('./img/icon_app.png')
    sizeImg = 50
    logo = a.resize((sizeImg, sizeImg))

    back_im = im_size.copy()

    back_im.paste(
        logo, (int(((back_im.size[0] - sizeImg))/2), int(((back_im.size[1] - sizeImg))/2)))
    print(back_im.size[0], back_im.size[0])

    back_im.save('./QR/' + name + '.png', quality=0)


# for root, dirs, files in os.walk(your_path, topdown=False):
#     for name in files:
#         # print(os.path.join(name))

#         r.add_picture('./QR/' + name)
#         document.save('QRAwaco.docx')

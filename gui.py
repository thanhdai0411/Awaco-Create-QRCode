from tkinter import *
from tkinter import ttk
from tkinter import messagebox
import requests

from time import sleep
import sys

from docx import Document
from docx.shared import Inches
from docx.shared import Mm
import os
from time import sleep
import base64

import urllib.request
import csv


import pyqrcode
import png
from pyqrcode import QRCode
from PIL import Image, ImageTk

document = Document()
p = document.add_paragraph()
r = p.add_run()

# set width page docx
section = document.sections[0]
# A3
section.page_height = Mm(120) #420  #280
section.page_width = Mm(460)   #297  #65

section.left_margin = Mm(5)
section.right_margin = Mm(5)
section.top_margin = Mm(5)
section.bottom_margin = Mm(5)

section.header_distance = Mm(5)
section.footer_distance = Mm(5)


qrData = []
msData = []



win= Tk()
win.title('Generate QRCode by NTĐ')
win.geometry("335x255") # x , y
win.resizable(0,0)

# url = "https://res.cloudinary.com/image-awaco/image/upload/v1675994608/utils/icon_app_l7g7ii.png"

# im = Image.open(requests.get(url, stream=True).raw)

# photo = ImageTk.PhotoImage(im)
# win.wm_iconphoto(False, photo)



def get_value():
    
    prefix = entry_prefix.get()
    code_prefix = entry_code_prefix.get()
    number_qr = entry_number.get()
    number_special = code_special.get()
    e_link = entry_link.get()
    type_export = variable.get()

    # print(number_special)
    # return

    print(type_export)
    ask = messagebox.askquestion('QR Code', 'Bạn chắc chắn tạo mã với tiền tố ' + prefix + ' số lương ' + number_qr + ' link khi quét bằng Camera thường ' + e_link + ' in thành file ' + type_export)
    if(ask =='no'):
        return 
    
    

    if(type_export == "WORD-PNG") :

        print(" >>>> Running! Waiting...")

        directory = prefix
        currentFolder = os.getcwd()
        parent_dir = currentFolder
        path = os.path.join(parent_dir, directory)
        os.mkdir(path)

        

        for index in range(int(number_qr)) :
            
            hash_prefix = base64.b64encode(prefix.encode('utf-8')).decode('utf-8')
            hash_index = base64.b64encode(str(index).encode('utf-8')).decode('utf-8')
            
            link = ""
            if number_special :
                link = e_link + "/?EVENT_RANDOM="
            else :
                link = e_link + "/?n="


            a = hash_prefix.replace('=','ntd')
            b = hash_index.replace('=','nhn')

            # code = link + b + 'NTD' + a 

            encode_number = b + 'NTD' + a

            mybytes = encode_number.encode()

            qrMs = int.from_bytes(mybytes, 'big')

            if number_special :
                
                first = str(code_prefix) # 2

                third = str(qrMs)[-5:-1] # 4

                summ = int(first) * int(third)

                four = str(summ)[0:2]

                ms =   third + first + four + number_special    # 4 + 2 + 2 + 9999

                qr_code_ = link  + third + first + four + number_special 


            else :
                second = str(qrMs)[4:8] # 4
                first = str(code_prefix) # 2
                third = str(qrMs)[-5:-1] # 4

                summ = int(first) + int(second)  + int(third)


                four = str(summ)[0:2]


                ms =  second + first + third + four
                qr_code_ = link  + ms



            qrData.append(qr_code_)
            msData.append(ms)


        for qr in qrData:
            
            c = qr.split('?')
            count = 0 
            arr= c[0]

            for index in arr :
                count = count + 1

            path_save = path + '\\' + qr[count+1:] + '.png'

            url = pyqrcode.create(qr,error='H')
            url.png(path_save, scale = 3,quiet_zone=0)

            im = Image.open(path_save)
            sizeImg = 140

            qrCodeImg = im.resize((sizeImg,sizeImg))

            qrCodeImg = qrCodeImg.convert("RGBA")
            # print(qrCodeImg.size[0],qrCodeImg.size[0] )

            # print('123123')
            # urllib.request.urlretrieve("https://res.cloudinary.com/image-awaco/image/upload/v1659334918/qr_code/khung_v1_x4xlng.png","khung.png")
            # khung = Image.open("D:\Workspace\CT\AppGenerateQR\img\khung_v1.png")
            # khung = Image.open("khung.png")

            
            url = "https://res.cloudinary.com/image-awaco/image/upload/v1659334918/qr_code/khung_v1_x4xlng.png"
            khung = Image.open(requests.get(url, stream=True).raw)



            back_im = khung.copy()

            back_im.paste(qrCodeImg, (55,27))
            # print(back_im.size[0],back_im.size[0] )

            back_im.save(path_save, quality=0)


        for root, dirs, files in os.walk(path, topdown=False):
            for name in files:
                r.add_picture(path + '\\' + name)
                document.save(prefix + '.docx')
    else:
        
        print(" >>>> Running! Waiting...")
        
        

        for index in range(int(number_qr)) :
            
            hash_prefix = base64.b64encode(prefix.encode('utf-8')).decode('utf-8')
            hash_index = base64.b64encode(str(index).encode('utf-8')).decode('utf-8')

            link = ""
            if number_special :
                link = e_link + "/?EVENT_RANDOM="
            else :
                link = e_link + "/?n="
        

            a = hash_prefix.replace('=','ntd')
            b = hash_index.replace('=','nhn')

            code = link + b + 'NTD' + a 
            
            encode_number = b + 'NTD' + a

            mybytes = encode_number.encode()

            qrMs = int.from_bytes(mybytes, 'big')

            if number_special :
                
                first = str(code_prefix) # 2

                third = str(qrMs)[-5:-1] # 4

                summ = int(first) * int(third)

                four = str(summ)[0:2]

                # ms =  third + first + four + number_special    # 4 + 2 + 2 + 9999
                ms =   third + first + four + number_special    # 4 + 2 + 2 + 9999


                qr_code_ = link  + third + first + four + number_special 
            else:


                second = str(qrMs)[4:8] # 4
                first = str(code_prefix) # 2
                third = str(qrMs)[-5:-1] # 4

                summ = int(first) + int(second)  + int(third)


                four = str(summ)[0:2]

                ms =  second + first + third + four
                qr_code_ = link  + ms



            qrData.append(qr_code_)
            msData.append(str(ms))


        with open(prefix+".csv", 'w',newline='') as csvfile: 
            csvwriter = csv.writer(csvfile,delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL) 
                
            csvwriter.writerow(["STT", "QR_CODE", "MS"]) 

            for index in range(len(qrData)):
                # for index in range(len(msData)):
                csvwriter.writerow([index+1,qrData[index], "'" + str(msData[index])]) 

    print(' >>>> Generate QR code success!')
    Label(win, text="Generate success!", font= ('Arial', 10),fg='#f00').grid(row = 5, column = 0, pady = 5, padx = 5)
    messagebox.showinfo('Thành công', 'Bạn đã tạo thành công ' + number_qr + ' mã QR code')



# prefix
l1=Label(win, text="Nhập tiền tố", font= ('Arial', 13))
l1.grid(row = 0, column = 0, pady = 5, padx = 5)

entry_prefix= ttk.Entry(win,font=('Arial', 10),width=25)
entry_prefix.grid(row = 0, column = 1, pady = 5)

# code prefix
l1_2=Label(win, text="Nhập mã tiền tố", font= ('Arial', 13))
l1_2.grid(row = 1, column = 0, pady = 5, padx = 5)

entry_code_prefix= ttk.Entry(win,font=('Arial', 10),width=25)
entry_code_prefix.grid(row = 1, column = 1, pady = 5)

# number count
l2 = Label(win, text="Nhập số lượng", font= ('Arial', 13))
l2.grid(row = 2, column = 0, pady = 5, padx = 5)

entry_number= ttk.Entry(win,font=('Arial', 10),width=25)
entry_number.grid(row = 2, column = 1, pady = 5)


# 4 number special
l2 = Label(win, text="Mã đặc biệt", font= ('Arial', 13))
l2.grid(row = 3, column = 0, pady = 5, padx = 5)

code_special= ttk.Entry(win,font=('Arial', 10),width=25)
code_special.grid(row = 3, column = 1, pady = 5)


# link redirect
l3 = Label(win, text="Nhập liên kết", font= ('Arial', 13))
l3.grid(row = 4, column = 0, pady = 5, padx = 5)

v = StringVar(win)
v.set("https://app.awaco.vn")

entry_link= ttk.Entry(win,font=('Arial', 10),width=25, textvariable=v)
entry_link.grid(row = 4, column = 1, pady = 5)

# choose type export

l4 = Label(win, text="Chọn loại file", font= ('Arial', 13))
l4.grid(row = 5, column = 0, pady = 5, padx = 5)

variable = StringVar(win)
variable.set("CSV") 

type_ = OptionMenu(win, variable , "WORD-PNG", "CSV")
type_.grid(row = 5, column = 1, pady = 5)

# entry_type= ttk.Entry(win,font=('Arial', 10),width=10, textvariable=variable)
# entry_type.grid(row = 3, column = 1, pady = 5)

# button
button= ttk.Button(win, text="Xác nhận", command= get_value)
button.grid(row = 6, column = 1, pady = 5)



win.mainloop()

#pyinstaller.exe --onefile --icon=icon_app.png gui.py
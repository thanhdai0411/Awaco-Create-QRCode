from docx import Document
from docx.shared import Inches
from docx.shared import Mm
import os
from time import sleep
import base64


import urllib.request
import csv

import pyqrcode
import png
from pyqrcode import QRCode
from PIL import Image

document = Document()
p = document.add_paragraph()
r = p.add_run()

# set width page docx
section = document.sections[0]
# A3
section.page_height = Mm(120) #420  #280
section.page_width = Mm(460)   #297  #65

section.left_margin = Mm(5)
section.right_margin = Mm(5)
section.top_margin = Mm(5)
section.bottom_margin = Mm(5)

section.header_distance = Mm(5)
section.footer_distance = Mm(5)


qrData = []

print('Nhập số lượng QR code muốn in: ')
number_qr = input()

print('Nhập tiền tố bắt đầu: ')
prefix = input()



for index in range(int(number_qr)) :
    
    hash_prefix = base64.b64encode(prefix.encode('utf-8')).decode('utf-8')
    hash_index = base64.b64encode(str(index).encode('utf-8')).decode('utf-8')

    e_link = "https://awaco.herokuapp.com"

    link = e_link + "/?n="

    a = hash_prefix.replace('=','ntd')
    b = hash_index.replace('=','nhn')
    code = link + b + 'NTD' + a 
    qrData.append(code)

print(qrData)

# for qr in qrData:
            
#     c = qr.split('?')5

#     count = 0 
#     arr= c[0]

#     for index in arr :
#         count = count + 1

#     path_save = path + '\\' + qr[count+1:] + '.png'

#     url = pyqrcode.create(qr,error='H')
#     url.png(path_save, scale = 3,quiet_zone=0)

#     im = Image.open(path_save)
#     sizeImg = 140

#     qrCodeImg = im.resize((sizeImg,sizeImg))

#     qrCodeImg = qrCodeImg.convert("RGBA")
#     # print(qrCodeImg.size[0],qrCodeImg.size[0] )

    
#     urllib.request.urlretrieve("https://res.cloudinary.com/image-awaco/image/upload/v1659334918/qr_code/khung_v1_x4xlng.png","khung.png")
#     khung = Image.open("khung.png")

#     back_im = khung.copy()

#     back_im.paste(qrCodeImg, (55,27))
#     # print(back_im.size[0],back_im.size[0] )

#     back_im.save(path_save, quality=0)


# for root, dirs, files in os.walk(path, topdown=False):
#     for name in files:
#         r.add_picture(path + '\\' + name)
#         document.save(prefix + '.docx')
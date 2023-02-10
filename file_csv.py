from docx import Document
from docx.shared import Inches
from docx.shared import Mm
import os
from time import sleep
import base64

import csv
import urllib.request


import pyqrcode
import png
from pyqrcode import QRCode
from PIL import Image

document = Document()
p = document.add_paragraph()
r = p.add_run()

# set width page docx
section = document.sections[0]
# A3
section.page_height = Mm(120) #420  #280
section.page_width = Mm(460)   #297  #65

section.left_margin = Mm(5)
section.right_margin = Mm(5)
section.top_margin = Mm(5)
section.bottom_margin = Mm(5)

section.header_distance = Mm(5)
section.footer_distance = Mm(5)



print('Nhập số lượng QR code muốn in: ')
number_qr = input()

print('Nhập tiền tố bắt đầu: ')
prefix = input()


directory = prefix
currentFolder = os.getcwd()
parent_dir = currentFolder
path = os.path.join(parent_dir, directory)
# os.mkdir(path)


qrData = []
msData = []



for index in range(int(number_qr)) :
    
    hash_prefix = base64.b64encode(prefix.encode('utf-8')).decode('utf-8')
    hash_index = base64.b64encode(str(index).encode('utf-8')).decode('utf-8')

    s = "https://app.awaco.vn/?"

    a = hash_prefix.replace('=','ntd')
    b = hash_index.replace('=','nhn')

    c = b + 'NTD' + a

    print(c)


    mybytes = c.encode()

    qrMs = int.from_bytes(mybytes, 'big')

    print(qrMs)
    
    first = str(qrMs)[4:8]
    second = str(qrMs)[12:16]
    third = str(qrMs)[-5:-1]

    print(first)
    print(second)
    print(third)

    ms = first + second + third
    code = s + c 


    print('=============================')

    print(ms)
    print(code)

    # print(code)

    print('=============================')
    
    
    msData.append(ms)
    qrData.append(code)

print(qrData)
print(msData)